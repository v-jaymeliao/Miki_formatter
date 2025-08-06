from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def highlight_table_recursive(tbl, colors, color_idx):
    color = colors[color_idx[0] % len(colors)]
    color_idx[0] += 1  # 每進入一個表格就+1
    for row in tbl.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
                    nsdecls('w'), color
                ))
            )
            for nested_tbl in cell.tables:
                highlight_table_recursive(nested_tbl, colors, color_idx)

def highlight_table(doc_path):
    doc = Document(doc_path)
    colors = ["FFFF00", "00FF00", "00B0F0", "FFC000", "FF99CC"]  # 黃、綠、藍、橘、粉紅
    color_idx = [0]  # 用 list 包起來，讓遞迴時可以修改

    for tbl in doc.tables:
        highlight_table_recursive(tbl, colors, color_idx)
    
    doc.save('highlighted_' + doc_path)
    print("所有表格（包含巢狀）都已經標記不同顏色！")

# 測試使用
highlight_table('Support Services Customer Paginated Report.docx')
