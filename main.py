from docx import Document
from datetime import timedelta
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import glob
import argparse
import sys

# 嘗試導入 PDF 轉換模組
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: docx2pdf not available. Only Word files will be generated.")
    print("To enable PDF conversion, run: pip install docx2pdf")

# Function to convert time format [h]:mm to minutes
def time_to_minutes(time_str):
    if not time_str:
        return 0
    hours, minutes = map(int, time_str.split(':'))
    return hours * 60 + minutes

# Function to convert minutes back to [h]:mm format
def minutes_to_time(minutes):
    hours = minutes // 60
    minutes = minutes % 60
    return f"{hours}:{minutes:02d}"

# Function to sum time values in the 'Used' column
def sum_times(time_list):
    total_minutes = sum(time_to_minutes(time) for time in time_list)
    return minutes_to_time(total_minutes)

# Function to create or modify the table in Word
def edit_word_table(doc_path):
    doc = Document(doc_path)
    table = None

    # Search for the table with specific columns
    for tbl in doc.tables:
        headers = [cell.text.strip() for cell in tbl.rows[0].cells]
        
        if 'Package' in headers and 'Service' in headers and 'Type' in headers:
            table = tbl
            break

    if not table:
        print("Table not found.")
        return

    # Extract column indices for easier access
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    package_idx = headers.index('Package')
    service_idx = headers.index('Service')
    type_idx = headers.index('Type')
    purchased_idx = headers.index('Purchased')
    used_idx = headers.index('Used')
    remaining_idx = headers.index('Remaining')
    trend_idx = headers.index('Trend')

    # Get the current rows data for sum calculations
    purchased_values = []
    used_values = []
    for row in table.rows[1:]:  # Skip header row
        purchased_values.append(row.cells[purchased_idx].text.strip())
        used_values.append(row.cells[used_idx].text.strip())

    # Sum the 'Purchased' and 'Used' values
    total_purchased = sum(float(value) for value in purchased_values if value)
    total_used = sum_times(used_values)

    # Add a new row for the 'Total'
    new_row = table.add_row()
    new_row.cells[service_idx].text = 'Total'
    new_row.cells[purchased_idx].text = str(total_purchased)
    new_row.cells[used_idx].text = total_used
    new_row.cells[remaining_idx].text = ''
    new_row.cells[trend_idx].text = ''

    # Save the modified document
    doc.save('modified_' + doc_path)




def highlight_table_yellow(tbl):
    # Test: Mark with color first
    # for row in tbl.rows:
    #     for cell in row.cells:
    #         cell._element.get_or_add_tcPr().append(
    #             parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
    #         )
    # Get column indices
    headers = [cell.text.strip() for cell in tbl.rows[0].cells]
    col_count = len(headers)
    # Check if the last row's Service column already has 'Total'
    service_idx = None
    for idx, h in enumerate(headers):
        if h == 'Service':
            service_idx = idx
            break
    if service_idx is not None and tbl.rows[-1].cells[service_idx].text.strip() == 'Total':
        return  # Already has Total row, skip adding

    # Prepare for summation
    purchased_sum = 0
    used_sum = 0
    remaining_sum = 0
    purchased_idx = used_idx = remaining_idx = trend_idx = service_idx = None
    for idx, h in enumerate(headers):
        if h == 'Purchased':
            purchased_idx = idx
        if h == 'Used':
            used_idx = idx
        if h == 'Remaining':
            remaining_idx = idx
        if h == 'Trend':
            trend_idx = idx
        if h == 'Service':
            service_idx = idx
    # Sum up purchased, used, remaining
    for row in tbl.rows[1:]:
        if purchased_idx is not None:
            val = row.cells[purchased_idx].text.strip()
            if val:
                try:
                    purchased_sum += int(val)
                except:
                    pass
        if used_idx is not None:
            val = row.cells[used_idx].text.strip()
            if val:
                try:
                    h, m = map(int, val.split(':'))
                    used_sum += h*60 + m
                except:
                    pass
        if remaining_idx is not None:
            val = row.cells[remaining_idx].text.strip()
            if val:
                try:
                    h, m = map(int, val.split(':'))
                    remaining_sum += h*60 + m
                except:
                    pass
    # Add a new row, fill only specified columns
    new_row = tbl.add_row()
    
    # Get the previous row's format as reference
    last_data_row = tbl.rows[-2]  # The last row before the new one
    
    for idx in range(col_count):
        header = headers[idx]
        cell = new_row.cells[idx]
        reference_cell = last_data_row.cells[idx]
        
        # Copy format properties
        if reference_cell._element.find('.//w:rPr', reference_cell._element.nsmap) is not None:
            # Copy font format
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if reference_cell.paragraphs and reference_cell.paragraphs[0].runs:
                        ref_run = reference_cell.paragraphs[0].runs[0]
                        if ref_run.font.name:
                            run.font.name = ref_run.font.name
                        if ref_run.font.size:
                            run.font.size = ref_run.font.size
                        if ref_run.bold is not None:
                            run.bold = ref_run.bold
                        if ref_run.italic is not None:
                            run.italic = ref_run.italic
        
        # Copy paragraph alignment
        if reference_cell.paragraphs and cell.paragraphs:
            if reference_cell.paragraphs[0].alignment is not None:
                cell.paragraphs[0].alignment = reference_cell.paragraphs[0].alignment
        
        # Set content
        content = ''
        if header == 'Service':
            content = ' Total'
        elif header == 'Type':
            content = ' Hour'
        elif header == 'Purchased':
            content = f"{purchased_sum}"
        elif header == 'Used':
            h = used_sum // 60
            m = used_sum % 60
            content = f"{h}:{m:02d}"
        elif header == 'Remaining':
            h = remaining_sum // 60
            m = remaining_sum % 60
            content = f"{h}:{m:02d}"
        elif header == 'Trend':
            percent = 0.0
            if purchased_sum > 0:
                percent = (used_sum / 60)  / purchased_sum * 100 
            content = f"{percent:.2f}%"
        else:
            content = ''
        
        # Clear existing content and recreate
        cell.text = ''
        paragraph = cell.paragraphs[0]
        
        # Copy alignment
        if reference_cell.paragraphs and reference_cell.paragraphs[0].alignment is not None:
            paragraph.alignment = reference_cell.paragraphs[0].alignment
        
        # Add new run and set format
        run = paragraph.add_run(content)
        
        # Copy font format
        if reference_cell.paragraphs and reference_cell.paragraphs[0].runs:
            ref_run = reference_cell.paragraphs[0].runs[0]
            if ref_run.font.name:
                run.font.name = ref_run.font.name
            if ref_run.font.size:
                run.font.size = ref_run.font.size
            if ref_run.bold is not None:
                run.bold = ref_run.bold
            if ref_run.italic is not None:
                run.italic = ref_run.italic
        
        # Add yellow background only to cells with content (from "Total" onwards)
        if content.strip():  # Only apply yellow background if cell has content
            cell._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
            )

def format_and_calc_table(doc_path):
    doc = Document(doc_path)
    target_headers = ["Package", "Service", "Type", "Purchased", "Used", "Remaining", "Trend"]

    def recursive_search(tbls):
        for tbl in tbls:
            if len(tbl.rows) > 0:
                headers = [cell.text.strip() for cell in tbl.rows[0].cells]
                if headers == target_headers:
                    highlight_table_yellow(tbl)
            # Recursively search nested tables
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.tables:
                        recursive_search(cell.tables)
    recursive_search(doc.tables)
    
    # Create success subdirectories in the same directory as the original file
    dir_path = os.path.dirname(doc_path)
    success_docx_dir = os.path.join(dir_path, "success_docx")
    success_pdf_dir = os.path.join(dir_path, "success_pdf")
    
    # Ensure success directories exist
    if not os.path.exists(success_docx_dir):
        os.makedirs(success_docx_dir)
        print(f"Created directory: {success_docx_dir}")
    
    if not os.path.exists(success_pdf_dir):
        os.makedirs(success_pdf_dir)
        print(f"Created directory: {success_pdf_dir}")
    
    filename = os.path.basename(doc_path)
    name, ext = os.path.splitext(filename)
    
    # Save Word file
    docx_outname = os.path.join(success_docx_dir, f"{name}{ext}")
    doc.save(docx_outname)
    print(f"Formatted and saved Word file as {docx_outname}")
    
    # Convert to PDF (if available)
    pdf_outname = os.path.join(success_pdf_dir, f"{name}.pdf")
    if PDF_AVAILABLE:
        try:
            # 檢查 Word 文件是否可讀取
            if not os.path.exists(docx_outname):
                raise FileNotFoundError(f"Word file not found: {docx_outname}")
            
            # 等待一下確保文件完全寫入磁盤
            import time
            time.sleep(0.5)
            
            # 嘗試 PDF 轉換，使用絕對路徑
            docx_abs_path = os.path.abspath(docx_outname)
            pdf_abs_path = os.path.abspath(pdf_outname)
            
            # 確保目標目錄存在
            os.makedirs(os.path.dirname(pdf_abs_path), exist_ok=True)
            
            print(f"  Attempting PDF conversion...")
            convert(docx_abs_path, pdf_abs_path)
            
            # Verify PDF was created and force memory cleanup
            if os.path.exists(pdf_abs_path) and os.path.getsize(pdf_abs_path) > 0:
                print(f"✓ Converted to PDF as {pdf_outname}")
                # Force cleanup of PDF conversion resources
                import gc
                gc.collect()
            else:
                raise Exception("PDF file was not created or is empty")
                
        except Exception as e:
            error_msg = str(e)
            print(f"Warning: PDF conversion failed for {docx_outname}")
            print(f"  Reason: {error_msg}")
            
            # 提供更具體的建議
            if "'NoneType' object has no attribute" in error_msg:
                print(f"  This usually means Microsoft Word is not properly installed or accessible.")
                print(f"  Try: 1) Ensure Microsoft Word is installed and licensed")
                print(f"       2) Try running as administrator")
                print(f"       3) Close any open Word documents")
            elif "Access is denied" in error_msg:
                print(f"  File access denied. Try running as administrator.")
            else:
                print(f"  You can manually open the Word file and save it as PDF")
            pdf_outname = None
    else:
        print(f"PDF conversion skipped (docx2pdf not available)")
        pdf_outname = None
    
    return docx_outname, pdf_outname

def batch_process_documents(input_path, recursive=True, file_pattern="*.docx"):
    """
    Batch process Word documents
    
    Args:
        input_path: Can be a single file path or directory path
        recursive: Whether to recursively search subdirectories
        file_pattern: File filter pattern, default is *.docx
    """
    processed_files = []
    failed_files = []
    
    # Check if input path is a file or directory
    if os.path.isfile(input_path):
        # Single file processing
        if input_path.lower().endswith('.docx'):
            try:
                print(f"Processing file: {input_path}")
                docx_output, pdf_output = format_and_calc_table(input_path)
                processed_files.append({'docx': docx_output, 'pdf': pdf_output})
                print(f"✓ Successfully processed: {input_path}")
            except Exception as e:
                print(f"✗ Processing failed: {input_path} - Error: {str(e)}")
                failed_files.append(input_path)
        else:
            print(f"Skipping non-Word file: {input_path}")
    
    elif os.path.isdir(input_path):
        # Directory processing
        if recursive:
            # Recursively search all subdirectories
            pattern = os.path.join(input_path, "**", file_pattern)
            docx_files = glob.glob(pattern, recursive=True)
        else:
            # Search current directory only
            pattern = os.path.join(input_path, file_pattern)
            docx_files = glob.glob(pattern)
        
        print(f"Found {len(docx_files)} Word files")
        
        # Memory management: Process files in batches
        batch_size = 5  # Process 5 files at a time to manage memory
        total_files = len(docx_files)
        
        for i in range(0, total_files, batch_size):
            batch_files = docx_files[i:i+batch_size]
            batch_num = (i // batch_size) + 1
            total_batches = (total_files + batch_size - 1) // batch_size
            
            print(f"\n--- Processing Batch {batch_num}/{total_batches} ({len(batch_files)} files) ---")
            
            for docx_file in batch_files:
                # Skip already formatted files
                if os.path.basename(docx_file).startswith('unused_'): # if you want to skip files that have been processed
                    print(f"Skipping already formatted file: {docx_file}")
                    continue
                    
                try:
                    print(f"Processing file ({i + batch_files.index(docx_file) + 1}/{total_files}): {os.path.basename(docx_file)}")
                    docx_output, pdf_output = format_and_calc_table(docx_file)
                    processed_files.append({'docx': docx_output, 'pdf': pdf_output})
                    print(f"✓ Successfully processed: {os.path.basename(docx_file)}")
                    
                    # Show completion summary for this file
                    pdf_status = "✓ PDF generated" if pdf_output else "⚠ PDF failed"
                    print(f"  → Word: ✓ | PDF: {pdf_status}")
                    
                    # Memory cleanup after each file
                    import gc
                    gc.collect()
                    
                except Exception as e:
                    print(f"✗ Processing failed: {os.path.basename(docx_file)} - Error: {str(e)}")
                    failed_files.append(docx_file)
            
            # Memory cleanup after each batch
            import gc
            gc.collect()
            
            # Show batch completion summary
            batch_successful = len([f for f in processed_files[-len(batch_files):] if f['docx']])
            batch_failed_count = len(batch_files) - batch_successful
            print(f"\n📊 Batch {batch_num} Summary:")
            print(f"   ✓ Successfully processed: {batch_successful} files")
            if batch_failed_count > 0:
                print(f"   ✗ Failed: {batch_failed_count} files")
            
            # Small pause between batches to allow system recovery
            if batch_num < total_batches:  # Don't pause after the last batch
                print(f"Batch {batch_num} completed. Pausing for memory cleanup...")
                import time
                time.sleep(1)  # 1 second pause
    
    else:
        print(f"Error: Path does not exist - {input_path}")
        return
    
    # Show processing results summary
    print("\n" + "="*50)
    print("🎉 Processing Results Summary:")
    print(f"Successfully processed: {len(processed_files)} files")
    print(f"Processing failed: {len(failed_files)} files")
    
    # Calculate PDF success rate
    pdf_successful = len([f for f in processed_files if f['pdf']])
    if processed_files:
        pdf_success_rate = (pdf_successful / len(processed_files)) * 100
        print(f"PDF conversion success rate: {pdf_success_rate:.1f}% ({pdf_successful}/{len(processed_files)})")
    
    if processed_files:
        print("\n📁 Successfully processed files:")
        for i, file_pair in enumerate(processed_files, 1):
            filename = os.path.basename(file_pair['docx'])
            pdf_icon = "📄" if file_pair['pdf'] else "❌"
            print(f"  {i:2d}. {filename}")
            print(f"      Word: ✓ | PDF: {pdf_icon}")
    
    if failed_files:
        print(f"\n❌ Failed to process files ({len(failed_files)}):")
        for i, file in enumerate(failed_files, 1):
            print(f"  {i:2d}. {os.path.basename(file)}")

def main():
    """Main function - handle command line arguments"""
    parser = argparse.ArgumentParser(description='Batch format tables in Word documents')
    parser.add_argument('input', help='Input file or directory path')
    parser.add_argument('--no-recursive', action='store_true', 
                       help='Do not recursively search subdirectories (only effective when input is directory)')
    parser.add_argument('--pattern', default='*.docx', 
                       help='File filter pattern (default: *.docx)')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Error: Path does not exist - {args.input}")
        sys.exit(1)
    
    recursive = not args.no_recursive
    
    print("Miki Word Document Formatter")
    print("="*30)
    print(f"Input path: {args.input}")
    print(f"Recursive search: {'Yes' if recursive else 'No'}")
    print(f"File pattern: {args.pattern}")
    print()
    
    batch_process_documents(args.input, recursive, args.pattern)

if __name__ == "__main__":
    # If no command line arguments, provide interactive mode
    if len(sys.argv) == 1:
        print("Miki Word Document Formatter - Interactive Mode")
        print("="*40)
        
        while True:
            input_path = input("Please enter file or directory path (enter 'q' to exit): ").strip()
            
            if input_path.lower() == 'q':
                break
                
            if not os.path.exists(input_path):
                print(f"Error: Path does not exist - {input_path}")
                continue
            
            # If it's a directory, ask whether to search recursively
            recursive = True
            if os.path.isdir(input_path):
                choice = input("Search subdirectories recursively? (y/n, default y): ").strip().lower()
                if choice in ['n', 'no']:
                    recursive = False
            
            print()
            batch_process_documents(input_path, recursive)
            print("\n" + "="*50 + "\n")
    else:
        main()